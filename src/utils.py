import logging
import os
import sys
from typing import Dict, List
import openai

from google.cloud import storage
from docx import Document

from constants import MENTOR_BOT_PROMPT

OPENAI_API_KEY = os.getenv('OPENAI_API_KEY')

# Initialize OpenAI API
client = openai
client.api_key = OPENAI_API_KEY

logging.basicConfig(format="%(asctime)s,%(msecs)d %(levelname)-8s [%(filename)s:%(lineno)d] %(message)s",
                    datefmt="%Y-%m-%d:%H:%M:%S",
                    level=logging.INFO,
                    stream=sys.stdout)

conversation_history: List[Dict[str, str]] = []  # Store all exchanges

def query_chatgpt(prompt: str,
                  model="gpt-4o-mini") -> str:
    """
    Calls the OpenAI ChatGPT API with the given prompt and returns the response.

    Parameters:
    prompt (str): The prompt to send to the ChatGPT API.

    Returns:
    str: The response from the ChatGPT API.
    "gpt-4o-2024-05-13"
    """
    try:
        response = client.chat.completions.create(
            model=model,
            stream=False,
            messages=[
                {"role": "system", "content": MENTOR_BOT_PROMPT},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7
        )
        # logging.info(f"response: {response} with type {type(response)}")
    except Exception as e:
        logging.error(f"Failed to query ChatGPT: {e}")    
    return response.choices[0].message.content.strip()


def create_prompt_business_plan(past_responses: List[Dict[str, str]]) -> str:
    """
    Creates a prompt for ChatGPT to determine the next question to ask the user
    based on their previous responses and the sections of a business plan.

    Parameters:
    past_responses (List[Dict[str, str]]): List containing user's previous responses.

    Returns:
    str: The formatted prompt for querying ChatGPT.
    """

    prompt = f"""
    You are an AI expert at collecting relevant information from entrepreneurs to create a business plan. 
    Based on the user's past responses, determine whether enough information has been provided for the business plan sections. 
    If so, simply output the string "3have_info3", and nothing else.
    If not, ask a detailed question to collect information on another section that has not been covered yet.

    Past responses:
    {past_responses}

    Ensure that you ask specific questions and try to minimize follow-up questions.
    """

    return prompt


def format_conversation_history(conversation_history):
    formatted_history = ""
    for message in conversation_history:
        role = message['role']
        content = message['content']
        formatted_history += f"{role.capitalize()}: {content}\n"
    return formatted_history    


def create_prompt_mentor_bot(past_responses: List[Dict[str, str]],
                             current_question) -> str:
    """
    Creates a prompt for ChatGPT to determine the next question to ask the user
    based on their previous responses and the sections of a business plan.

    Parameters:
    past_responses (List[Dict[str, str]]): List containing user's previous responses.

    Returns:
    str: The formatted prompt for querying ChatGPT.
    """

    formatted_history = format_conversation_history(past_responses)

    prompt = (
    "You are MentorGPT, an Entrepreneurship Mentor with over 20 years of experience guiding entrepreneurs.\n"
    "Your expertise includes advising on entrepreneurial models, business concepts, financial accounting, scaling strategies, go-to-market strategies, marketing, and regional business contexts.\n"
    "Answer the following question from the user, taking into account the information provided in the conversation so far.\n"
    f"{current_question}"
    "Conversation history:\n"
    f"{formatted_history}\n"
    )

    return prompt

def update_conversation_history(user_input: str, assistant_response: str) -> None:
    """
    Updates the conversation history with the latest user input and assistant's response.

    Args:
        user_input (str): The input provided by the user.
        assistant_response (str): The response from the assistant (ChatGPT).
    """
    conversation_history.append({
        'role': 'user',
        'content': user_input
    })
    conversation_history.append({
        'role': 'assistant',
        'content': assistant_response
    })

def create_prompt_generate_business_plan(conversation_history: List[Dict[str, str]]) -> str:
    prompt = f"""
    You are an expert at creating a business plan. Based on the below list of previous responses
    from the user, generate a detailed business plan following the below structure. 

    "Business Description",
    "Business Mission",
    "Business goals",
    "Target Market and Industry",
    "Target Customer",


    Previous responses from user:
    {conversation_history}
    """
    return prompt


def upload_file_to_gcs(file_name: str, bucket_name: str, blob_name: str = None) -> str:
    """
    Uploads a file to Google Cloud Storage and returns the public URL.

    Parameters:
    - file_name (str): The local file to upload.
    - bucket_name (str): The GCS bucket name.
    - blob_name (str): The name of the object in GCS. If not specified, file_name will be used.

    Returns:
    - str: The public URL of the uploaded file.
    """
    try:
        # Initialize the Google Cloud Storage client
        storage_client = storage.Client()
        bucket = storage_client.bucket(bucket_name)
        blob = bucket.blob(blob_name or os.path.basename(file_name))

        # Upload the file to GCS
        blob.upload_from_filename(file_name)

        # # Make the blob publicly viewable
        # blob.make_public()

        # Return the public URL
        return blob.public_url
    except Exception as e:
        logging.error(f"Failed to upload {file_name} to GCS: {e}")
        return None
    
def generate_business_plan_document(business_plan: str, file_format: str = 'docx') -> str:
    """
    Generate a business plan document and save it as a file.

    Parameters:
    - business_plan (str): The business plan content.
    - file_format (str): The format to save the file in ('txt' or 'docx').

    Returns:
    - str: The path to the saved file.
    """
    file_name = f"business_plan_agricollect.{file_format}"

    if file_format == 'docx':
        # Create a Word document
        doc = Document()
        doc.add_heading('AgriCollect Business Plan', 0)
        doc.add_paragraph(business_plan)

        # Save the document
        doc.save(file_name)
    else:
        # Save as plain text
        with open(file_name, 'w') as file:
            file.write(business_plan)

    return file_name


def split_message(message, max_length=1000):
    """
    Splits a message into chunks of a specified maximum length. Adds "..." to the end of each chunk 
    (except the last) to indicate continuation.

    Args:
        message (str): The message to be split.
        max_length (int): The maximum allowed length for each chunk (default is 1000 characters).

    Returns:
        list of str: A list of message chunks, each within the character limit.
    """

    # Split message into chunks of `max_length`
    chunks = [message[i:i + max_length] for i in range(0, len(message), max_length)]
    
    # Add "..." to all chunks except the last one to indicate continuation
    for i in range(len(chunks) - 1):
        chunks[i] += "..."
    
    return chunks



def query_chatgpt_assistant(question: str, assistant_id: str = "asst_xUtTiMSkYXGo0zidNvUpWJ38", thread_id: str = None) -> (str, str):
    """
    Query a custom assistant using OpenAI's client by creating a message in a thread,
    running the assistant, and retrieving the response. If no thread exists, a new one
    is created.

    Args:
        client (OpenAIClient): The OpenAI client instance.
        question (str): The question or prompt to send to the assistant.
        assistant_id (str): The unique identifier for the custom assistant created on OpenAI.
        thread_id (str): Optional existing thread ID for the conversation. If None, a new thread is created.

    Returns:
        tuple: (assistant's response, thread ID used in the interaction)

    Raises:
        Exception: If an error occurs during message creation, assistant run, or message retrieval.
    """
    try:
        # Create a new thread if one doesn't exist
        if thread_id is None:
            thread = client.beta.threads.create()
            thread_id = thread.id
            logging.info(f"New thread created with ID: {thread_id}")

        # Add the user's message to the thread
        client.beta.threads.messages.create(
            thread_id=thread_id,
            role="user",
            content=question
        )

        # Run the assistant on the specified thread
        run = client.beta.threads.runs.create(
            thread_id=thread_id,
            assistant_id=assistant_id
        )

        # Poll the run status until the assistant completes processing
        while True:
            run_status = client.beta.threads.runs.retrieve(
                thread_id=thread_id,
                run_id=run.id
            )
            
            if run_status.status == 'completed':
                # Retrieve assistant's response
                messages = client.beta.threads.messages.list(
                    thread_id=thread_id
                )

                assistant_response = ""
                for msg in messages.data:
                    if msg.role == "assistant":
                        content = msg.content[0].text.value
                        assistant_response = f"{content}\n"
                        break

                return assistant_response, thread_id

    except Exception as e:
        logging.error(f"Error querying the assistant: {e}")
        raise